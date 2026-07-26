from __future__ import annotations

import json
import shutil
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from reportlab.lib.colors import Color, HexColor, white
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
PACK_ROOT = ROOT / "brand-pack" / "DeepBridge Advisory Brand Pack"
LOGOS_DIR = PACK_ROOT / "01 Logos"
LETTERHEAD_DIR = PACK_ROOT / "02 Letterhead"
CONTRACTS_DIR = PACK_ROOT / "03 Contract Templates"
ONBOARDING_DIR = PACK_ROOT / "04 Onboarding"
DIGITAL_DIR = PACK_ROOT / "05 Digital"
GUIDE_DIR = PACK_ROOT / "06 Brand Guidelines"
PUBLIC_BRAND_DIR = ROOT / "public" / "brand"
ZIP_PATH = ROOT / "brand-pack" / "DeepBridge-Advisory-Brand-Pack.zip"

NAVY = "#071722"
NAVY_SOFT = "#142D38"
AQUA = "#43D0C7"
AQUA_SOFT = "#A9E9DF"
IVORY = "#F5F4EF"
WHITE = "#FFFFFF"
MUTED = "#56666C"
LINE = "#D7DDDC"
PALE_AQUA = "#E7F5F2"
CONTENT_DXA = 9638

ARIAL = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
ARIAL_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")
GEORGIA = Path("/System/Library/Fonts/Supplemental/Georgia.ttf")
GEORGIA_ITALIC = Path("/System/Library/Fonts/Supplemental/Georgia Italic.ttf")

COMPANY = {
    "brand": "DeepBridge Advisory",
    "legal_name": "DUSTDEEP LTD",
    "company_number": "16775578",
    "registered_in": "England and Wales",
    "registered_office": (
        "Kemp House, 152-160 City Road, London, United Kingdom, EC1V 2NX"
    ),
    "email": "hello@deepbridgeadvisory.co.uk",
    "website": "www.deepbridgeadvisory.co.uk",
    "linkedin": "linkedin.com/company/deepbridge-advisory",
}


def reset_output() -> None:
    if PACK_ROOT.exists():
        shutil.rmtree(PACK_ROOT)
    for directory in (
        LOGOS_DIR,
        LETTERHEAD_DIR,
        CONTRACTS_DIR,
        ONBOARDING_DIR,
        DIGITAL_DIR,
        GUIDE_DIR,
        PUBLIC_BRAND_DIR,
    ):
        directory.mkdir(parents=True, exist_ok=True)
    ZIP_PATH.parent.mkdir(parents=True, exist_ok=True)


def font(path: Path, size: int, index: int = 0) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(path), size=size, index=index)


def draw_tracked_text(
    draw: ImageDraw.ImageDraw,
    position: tuple[int, int],
    text: str,
    text_font: ImageFont.FreeTypeFont,
    fill: str,
    tracking: int,
) -> int:
    x, y = position
    for character in text:
        draw.text((x, y), character, font=text_font, fill=fill)
        box = draw.textbbox((x, y), character, font=text_font)
        x += box[2] - box[0] + tracking
    return x


def render_monogram(size: int, background: str | None = None) -> Image.Image:
    scale = 4
    canvas_size = size * scale
    image = Image.new(
        "RGBA",
        (canvas_size, canvas_size),
        background if background else (0, 0, 0, 0),
    )
    draw = ImageDraw.Draw(image)
    inset = int(canvas_size * 0.055)
    ring_width = max(4, int(canvas_size * 0.018))
    draw.ellipse(
        (inset, inset, canvas_size - inset, canvas_size - inset),
        fill=NAVY,
        outline="#4D8F91",
        width=ring_width,
    )
    draw.ellipse(
        (
            inset + ring_width * 2,
            inset + ring_width * 2,
            canvas_size - inset - ring_width * 2,
            canvas_size - inset - ring_width * 2,
        ),
        outline="#123342",
        width=max(2, ring_width // 2),
    )

    letter_font = font(GEORGIA, int(canvas_size * 0.49))
    baseline_y = int(canvas_size * 0.255)
    draw.text(
        (int(canvas_size * 0.205), baseline_y),
        "D",
        font=letter_font,
        fill=IVORY,
        stroke_width=max(1, int(canvas_size * 0.0015)),
        stroke_fill=IVORY,
    )
    draw.text(
        (int(canvas_size * 0.485), baseline_y),
        "B",
        font=letter_font,
        fill=AQUA,
        stroke_width=max(1, int(canvas_size * 0.0015)),
        stroke_fill=AQUA,
    )
    draw.line(
        (
            int(canvas_size * 0.535),
            int(canvas_size * 0.205),
            int(canvas_size * 0.455),
            int(canvas_size * 0.805),
        ),
        fill=AQUA,
        width=max(3, int(canvas_size * 0.013)),
    )
    return image.resize((size, size), Image.Resampling.LANCZOS)


def render_primary_logo(variant: str) -> Image.Image:
    width, height = 1800, 360
    image = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    mark = render_monogram(300)
    image.alpha_composite(mark, (24, 30))
    draw = ImageDraw.Draw(image)
    word_color = IVORY if variant == "light" else NAVY
    advisory_color = AQUA_SOFT if variant == "light" else "#176762"
    word_font = font(ARIAL_BOLD, 94)
    advisory_font = font(ARIAL_BOLD, 50)
    draw_tracked_text(
        draw,
        (388, 76),
        "DEEPBRIDGE",
        word_font,
        word_color,
        tracking=14,
    )
    draw_tracked_text(
        draw,
        (390, 205),
        "ADVISORY",
        advisory_font,
        advisory_color,
        tracking=24,
    )
    return image


def crop_transparent(image: Image.Image, padding: int = 12) -> Image.Image:
    bbox = image.getbbox()
    if bbox is None:
        return image
    left, top, right, bottom = bbox
    return image.crop(
        (
            max(0, left - padding),
            max(0, top - padding),
            min(image.width, right + padding),
            min(image.height, bottom + padding),
        )
    )


def make_logo_assets() -> dict[str, Path]:
    light = render_primary_logo("light")
    dark = render_primary_logo("dark")
    monogram = render_monogram(1024)

    assets = {
        "light": LOGOS_DIR / "DeepBridge-Advisory-Primary-Light.png",
        "dark": LOGOS_DIR / "DeepBridge-Advisory-Primary-Dark.png",
        "monogram": LOGOS_DIR / "DeepBridge-Advisory-Monogram-1024.png",
    }
    light.save(assets["light"], optimize=True)
    dark.save(assets["dark"], optimize=True)
    monogram.save(assets["monogram"], optimize=True)

    for size in (512, 256, 180, 64, 32):
        icon = render_monogram(size)
        icon.save(
            LOGOS_DIR / f"DeepBridge-Advisory-Monogram-{size}.png",
            optimize=True,
        )

    crop_transparent(light).save(
        PUBLIC_BRAND_DIR / "deepbridge-logo-light.png",
        optimize=True,
    )
    crop_transparent(dark).save(
        PUBLIC_BRAND_DIR / "deepbridge-logo-dark.png",
        optimize=True,
    )
    render_monogram(512).save(
        PUBLIC_BRAND_DIR / "deepbridge-monogram-512.png",
        optimize=True,
    )
    render_monogram(64).save(ROOT / "public" / "favicon-64.png", optimize=True)
    render_monogram(180).save(
        ROOT / "public" / "apple-touch-icon.png",
        optimize=True,
    )
    return assets


def set_cell_margins(
    cell,
    top: int = 100,
    start: int = 140,
    bottom: int = 100,
    end: int = 140,
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_border(cell, color: str = "D7DDDC", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color.replace("#", ""))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(
    run,
    name: str = "Arial",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY.replace("#", ""))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (18, NAVY, 14, 7),
        "Heading 2": (13, "#176762", 11, 5),
        "Heading 3": (11.5, NAVY_SOFT, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        list_style = styles[list_style_name]
        list_style.font.name = "Arial"
        list_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        list_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        list_style.font.size = Pt(10.5)
        list_style.paragraph_format.left_indent = Inches(0.5)
        list_style.paragraph_format.first_line_indent = Inches(-0.25)
        list_style.paragraph_format.space_after = Pt(6)
        list_style.paragraph_format.line_spacing = 1.15


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, size=7.5, color=MUTED)


def add_top_border(paragraph, color: str = "43D0C7", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), color.replace("#", ""))
    p_bdr.append(top)


def configure_section(section, logo_path: Path) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(22)
    section.right_margin = Mm(20)
    section.bottom_margin = Mm(22)
    section.left_margin = Mm(20)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(9)

    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_paragraph.add_run().add_picture(str(logo_path), width=Inches(2.15))

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Mm(170))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [8178, 1460], indent_dxa=0)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    for cell in (left, right):
        set_cell_margins(cell, top=60, start=0, bottom=0, end=0)
    legal = (
        f"{COMPANY['legal_name']} trading as {COMPANY['brand']} | "
        f"Company no. {COMPANY['company_number']} | "
        f"Registered office: {COMPANY['registered_office']} | "
        f"{COMPANY['email']}"
    )
    legal_paragraph = left.paragraphs[0]
    legal_paragraph.paragraph_format.space_after = Pt(0)
    legal_paragraph.paragraph_format.line_spacing = 1.0
    legal_run = legal_paragraph.add_run(legal)
    set_run_font(legal_run, size=7.2, color=MUTED)
    page_paragraph = right.paragraphs[0]
    add_page_number(page_paragraph)
    add_top_border(legal_paragraph)
    add_top_border(page_paragraph)


def new_document(logo_path: Path) -> Document:
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], logo_path)
    props = doc.core_properties
    props.title = "DeepBridge Advisory document template"
    props.subject = "DeepBridge Advisory branded business document"
    props.author = "DeepBridge Advisory"
    props.last_modified_by = "DeepBridge Advisory"
    props.keywords = "DeepBridge Advisory, DUSTDEEP LTD, template"
    return doc


def add_kicker(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=8.5, color="#176762", bold=True)


def add_title(
    doc: Document,
    title: str,
    subtitle: str | None = None,
    size: float = 28,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    set_run_font(run, name="Georgia", size=size, color=NAVY)
    if subtitle:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(20)
        sub_run = sub.add_run(subtitle)
        set_run_font(sub_run, size=11.5, color=MUTED)


def add_label_value_table(
    doc: Document,
    rows: list[tuple[str, str]],
    label_width: int = 2400,
    value_width: int = 7238,
) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [label_width, value_width])
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], PALE_AQUA)
        for cell in row.cells:
            set_cell_border(cell)
        label_p = row.cells[0].paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        label_run = label_p.add_run(label)
        set_run_font(label_run, size=8.5, color="#176762", bold=True)
        value_p = row.cells[1].paragraphs[0]
        value_p.paragraph_format.space_after = Pt(0)
        value_run = value_p.add_run(value)
        set_run_font(value_run, size=10, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_note_box(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_AQUA)
    set_cell_border(cell, color="A8D8D3")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label.upper()}  ")
    set_run_font(label_run, size=8.5, color="#176762", bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=9.5, color=NAVY_SOFT)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_checklist(
    doc: Document,
    items: list[tuple[str, str]],
    title: str | None = None,
) -> None:
    if title:
        doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [900, 3500, 5238])
    set_repeat_table_header(table.rows[0])
    headers = ("Status", "Item", "Notes / owner")
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, color=NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=8.5, color=WHITE, bold=True)
    for item, notes in items:
        row = table.add_row()
        values = ("[ ]", item, notes)
        for cell, value in zip(row.cells, values):
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=9.2, color=NAVY)
    set_table_geometry(table, [900, 3500, 5238])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_signature_block(doc: Document) -> None:
    doc.add_heading("Signatures", level=1)
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [4819, 4819])
    rows = (
        ("For DUSTDEEP LTD", "For [Client legal name]"),
        ("Name: [Insert]", "Name: [Insert]"),
        ("Title: [Insert]", "Title: [Insert]"),
        ("Signature / date: ____________________", "Signature / date: ____________________"),
    )
    for row_index, (row, values) in enumerate(zip(table.rows, rows)):
        for cell, value in zip(row.cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(value)
            set_run_font(
                run,
                size=9.5,
                color=NAVY,
                bold=row_index == 0,
            )


def make_letterhead(logo_path: Path) -> Path:
    doc = new_document(logo_path)
    add_kicker(doc, "Business correspondence")
    add_title(doc, "[Document title or subject]", "[Optional subtitle or reference]")
    add_label_value_table(
        doc,
        [
            ("Date", "[Day Month Year]"),
            ("To", "[Recipient name, title and organisation]"),
            ("Reference", "[Reference / project / agreement number]"),
        ],
    )
    greeting = doc.add_paragraph()
    greeting.add_run("[Dear Name],")
    for text in (
        "[Begin your correspondence here. Replace bracketed guidance with final text.]",
        "[Use short paragraphs and clear headings for longer letters. The DeepBridge "
        "header and legal footer will remain consistent on every page.]",
    ):
        doc.add_paragraph(text)
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(12)
    closing.add_run("Yours sincerely,")
    doc.add_paragraph("[Name]\n[Title]\nDeepBridge Advisory")
    path = LETTERHEAD_DIR / "DeepBridge-Advisory-Letterhead.docx"
    doc.save(path)
    return path


def make_contract_template(logo_path: Path) -> Path:
    doc = new_document(logo_path)
    add_kicker(doc, "Agreement template")
    add_title(
        doc,
        "[Agreement title]",
        "A branded document shell for counsel-approved contractual terms",
        size=31,
    )
    add_label_value_table(
        doc,
        [
            ("Parties", "DUSTDEEP LTD trading as DeepBridge Advisory and [Counterparty]"),
            ("Effective date", "[Day Month Year]"),
            ("Project / assignment", "[Reference]"),
            ("Document status", "[Draft / final / executed]"),
        ],
    )
    add_note_box(
        doc,
        "Important",
        "This file provides document structure and branding only. It is not a legal "
        "agreement and does not replace review by appropriately qualified legal, tax "
        "or employment-status advisers.",
    )
    for heading, body in (
        ("1. Parties and background", "[Insert counsel-approved wording.]"),
        ("2. Definitions and interpretation", "[Insert counsel-approved wording.]"),
        ("3. Scope and deliverables", "[Describe the agreed services, outputs and exclusions.]"),
        ("4. Commercial terms", "[Insert fees, invoicing, expenses and payment terms.]"),
        ("5. Responsibilities", "[Set out responsibilities for each party.]"),
        ("6. Confidentiality and data protection", "[Insert counsel-approved wording.]"),
        ("7. Compliance and engagement status", "[Insert assignment-specific wording.]"),
        ("8. Term, termination and consequences", "[Insert counsel-approved wording.]"),
        ("9. Liability and insurance", "[Insert counsel-approved wording.]"),
        ("10. General", "[Insert notices, assignment, variation and governing law.]"),
    ):
        doc.add_heading(heading, level=1)
        doc.add_paragraph(body)
    add_signature_block(doc)
    path = CONTRACTS_DIR / "DeepBridge-Advisory-Contract-Shell.docx"
    doc.save(path)
    return path


def make_consultant_onboarding(logo_path: Path) -> Path:
    doc = new_document(logo_path)
    add_kicker(doc, "Consultant enablement")
    add_title(
        doc,
        "Consultant onboarding",
        "A clear route from confirmation to a well-supported engagement",
        size=31,
    )
    add_label_value_table(
        doc,
        [
            ("Consultant", "[Full name]"),
            ("Client / programme", "[Organisation and programme]"),
            ("Assignment", "[Role / workstream]"),
            ("Start and duration", "[Date / expected duration]"),
            ("DeepBridge contact", "[Name / email / phone]"),
        ],
    )
    add_note_box(
        doc,
        "Secure information",
        "Do not send passport, identity, bank or tax documents through the public "
        "website. DeepBridge will confirm an appropriate secure route when documents "
        "are required.",
    )
    add_checklist(
        doc,
        [
            ("Assignment scope confirmed", "[Owner / date]"),
            ("Availability and working model confirmed", "[Owner / date]"),
            ("Commercial terms agreed", "[Owner / date]"),
            ("Consultant company and invoicing details reviewed", "[Owner / date]"),
            ("Right-to-work / work-authorisation review completed where applicable", "[Owner / date]"),
            ("Insurance evidence reviewed where required", "[Owner / date]"),
            ("Client onboarding steps completed", "[Owner / date]"),
            ("Access, equipment and first-day details issued", "[Owner / date]"),
        ],
        "Pre-engagement checklist",
    )
    doc.add_heading("Assignment snapshot", level=1)
    add_label_value_table(
        doc,
        [
            ("Primary outcome", "[What must this assignment deliver?]"),
            ("Key stakeholders", "[Names, roles and locations]"),
            ("Working model", "[Remote / hybrid / onsite, travel expectations]"),
            ("Reporting cadence", "[Meetings, status reporting and timesheets]"),
            ("Dependencies", "[Systems, teams, decisions and inputs]"),
            ("Escalation route", "[Client lead and DeepBridge contact]"),
        ],
    )
    doc.add_heading("First-week plan", level=1)
    for step in (
        "Confirm stakeholder introductions and programme context.",
        "Agree immediate priorities, expected outputs and decision points.",
        "Validate access, data and system requirements.",
        "Confirm reporting, time recording and communication routines.",
        "Raise any early delivery or engagement risks with DeepBridge.",
    ):
        doc.add_paragraph(step, style="List Number")
    doc.add_heading("During the engagement", level=1)
    for item in (
        "Maintain direct communication on progress, dependencies and risks.",
        "Tell DeepBridge promptly if scope, location or commercial expectations change.",
        "Use the agreed route for timesheets, expenses and invoicing.",
        "Handle client information in line with the assignment agreement and applicable policies.",
        "Keep DeepBridge informed of extension, transition or offboarding discussions.",
    ):
        doc.add_paragraph(item, style="List Bullet")
    add_note_box(
        doc,
        "Engagement status",
        "Joining the DeepBridge network or receiving this pack does not create an "
        "employment relationship, guarantee an assignment or override the applicable "
        "written agreement.",
    )
    path = ONBOARDING_DIR / "DeepBridge-Advisory-Consultant-Onboarding.docx"
    doc.save(path)
    return path


def make_client_kickoff(logo_path: Path) -> Path:
    doc = new_document(logo_path)
    add_kicker(doc, "Client delivery alignment")
    add_title(
        doc,
        "Programme kickoff",
        "A focused briefing template for specialist assignments and team build-outs",
        size=29,
    )
    add_label_value_table(
        doc,
        [
            ("Client", "[Organisation]"),
            ("Programme", "[Programme / workstream]"),
            ("DeepBridge lead", "[Name / email / phone]"),
            ("Kickoff date", "[Day Month Year]"),
            ("Document owner", "[Name and role]"),
        ],
    )
    doc.add_heading("What success looks like", level=1)
    doc.add_paragraph(
        "[State the programme outcome, the reason support is required now and the "
        "decisions or deliverables that will define success.]"
    )
    add_label_value_table(
        doc,
        [
            ("Required expertise", "[Role, capability and seniority]"),
            ("Deliverables", "[Specific outputs and milestones]"),
            ("Programme stage", "[Discovery / design / build / test / cutover / stabilisation]"),
            ("Location and travel", "[Working model and onsite expectations]"),
            ("Target start / duration", "[Date and expected term]"),
            ("Commercial parameters", "[Approved range / structure]"),
        ],
    )
    add_checklist(
        doc,
        [
            ("Scope and outcomes agreed", "[Owner / date]"),
            ("Stakeholder and interview panel confirmed", "[Owner / date]"),
            ("Working model and travel clarified", "[Owner / date]"),
            ("Commercial approval confirmed", "[Owner / date]"),
            ("Compliance and onboarding requirements identified", "[Owner / date]"),
            ("Interview and decision timeline agreed", "[Owner / date]"),
            ("Start-date dependencies identified", "[Owner / date]"),
        ],
        "Kickoff checklist",
    )
    doc.add_heading("Stakeholders and communication", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2400, 2300, 2400, 2538])
    set_repeat_table_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, ("Name", "Role", "Decision / input", "Cadence")):
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, color=NAVY)
        run = cell.paragraphs[0].add_run(value)
        set_run_font(run, size=8.5, color=WHITE, bold=True)
    for _ in range(4):
        row = table.add_row()
        for cell in row.cells:
            set_cell_border(cell)
            run = cell.paragraphs[0].add_run("[Insert]")
            set_run_font(run, size=9.2, color=MUTED)
    set_table_geometry(table, [2400, 2300, 2400, 2538])
    doc.add_heading("Risks, dependencies and open questions", level=1)
    add_label_value_table(
        doc,
        [
            ("Programme dependencies", "[Systems, data, teams, approvals]"),
            ("Delivery risks", "[Known constraints or uncertainties]"),
            ("Open questions", "[Items requiring resolution]"),
            ("Escalation route", "[Client sponsor and DeepBridge contact]"),
        ],
    )
    add_note_box(
        doc,
        "Advisory boundary",
        "Engagement structures and assignment-specific compliance requirements "
        "should be reviewed with appropriately qualified legal, tax, immigration "
        "or employment-status advisers where necessary.",
    )
    path = ONBOARDING_DIR / "DeepBridge-Advisory-Client-Kickoff.docx"
    doc.save(path)
    return path


def draw_pdf_tracked_text(
    pdf: canvas.Canvas,
    x: float,
    y: float,
    text: str,
    font_name: str,
    size: float,
    color: Color,
    tracking: float,
) -> None:
    pdf.setFont(font_name, size)
    pdf.setFillColor(color)
    cursor = x
    for character in text:
        pdf.drawString(cursor, y, character)
        cursor += stringWidth(character, font_name, size) + tracking


def make_brand_guide(assets: dict[str, Path]) -> Path:
    path = GUIDE_DIR / "DeepBridge-Advisory-Brand-Guide.pdf"
    width, height = A4
    pdf = canvas.Canvas(str(path), pagesize=A4)
    pdf.setTitle("DeepBridge Advisory Brand Guide")
    pdf.setAuthor("DeepBridge Advisory")

    # Cover
    pdf.setFillColor(HexColor(NAVY))
    pdf.rect(0, 0, width, height, fill=1, stroke=0)
    pdf.drawImage(
        ImageReader(str(assets["light"])),
        46,
        height - 150,
        width=300,
        height=60,
        mask="auto",
    )
    pdf.setFillColor(HexColor(AQUA_SOFT))
    pdf.setFont("Helvetica-Bold", 9)
    pdf.drawString(48, height - 230, "IDENTITY SYSTEM / VERSION 1.0")
    pdf.setFillColor(white)
    pdf.setFont("Times-Roman", 36)
    pdf.drawString(46, height - 305, "Clear expertise.")
    pdf.drawString(46, height - 348, "Considered delivery.")
    pdf.setFillColor(HexColor(AQUA_SOFT))
    pdf.setFont("Times-Italic", 26)
    pdf.drawString(46, height - 397, "One coherent DeepBridge identity.")
    pdf.setStrokeColor(HexColor(AQUA))
    pdf.setLineWidth(1)
    pdf.line(46, 92, width - 46, 92)
    pdf.setFillColor(HexColor("#8EA1A8"))
    pdf.setFont("Helvetica", 8)
    pdf.drawString(
        46,
        68,
        "DeepBridge Advisory | DUSTDEEP LTD | Company no. 16775578",
    )
    pdf.showPage()

    # Core identity
    pdf.setFillColor(HexColor(IVORY))
    pdf.rect(0, 0, width, height, fill=1, stroke=0)
    pdf.setFillColor(HexColor(NAVY))
    pdf.setFont("Times-Roman", 27)
    pdf.drawString(46, height - 72, "The DeepBridge identity")
    pdf.setFont("Helvetica", 10)
    text = pdf.beginText(46, height - 110)
    text.setLeading(15)
    text.textLines(
        "The identity is calm, precise and international. It should feel like a\n"
        "specialist advisory firm: confident without being loud, modern without\n"
        "looking like a software product or high-volume recruitment brand."
    )
    pdf.drawText(text)
    pdf.drawImage(
        ImageReader(str(assets["dark"])),
        46,
        height - 255,
        width=350,
        height=70,
        mask="auto",
    )
    pdf.setFont("Helvetica-Bold", 9)
    pdf.setFillColor(HexColor("#176762"))
    pdf.drawString(46, height - 300, "PRIMARY PALETTE")
    swatches = [
        ("Deep Navy", NAVY),
        ("Bridge Aqua", AQUA),
        ("Soft Aqua", AQUA_SOFT),
        ("Advisory Ivory", IVORY),
        ("Slate", MUTED),
    ]
    x = 46
    for name, color in swatches:
        pdf.setFillColor(HexColor(color))
        pdf.setStrokeColor(HexColor("#D3DAD9" if color == IVORY else color))
        pdf.roundRect(
            x,
            height - 390,
            88,
            58,
            4,
            fill=1,
            stroke=1 if color == IVORY else 0,
        )
        pdf.setFillColor(HexColor(NAVY))
        pdf.setFont("Helvetica-Bold", 7.5)
        pdf.drawString(x, height - 410, name)
        pdf.setFont("Helvetica", 7.5)
        pdf.drawString(x, height - 424, color)
        x += 100
    pdf.setFont("Helvetica-Bold", 9)
    pdf.setFillColor(HexColor("#176762"))
    pdf.drawString(46, height - 480, "TYPE SYSTEM")
    pdf.setFillColor(HexColor(NAVY))
    pdf.setFont("Times-Roman", 25)
    pdf.drawString(46, height - 525, "Georgia / editorial authority")
    pdf.setFont("Helvetica-Bold", 13)
    pdf.drawString(46, height - 560, "ARIAL BOLD / NAVIGATION AND LABELS")
    pdf.setFont("Helvetica", 10)
    pdf.drawString(46, height - 588, "Arial Regular / body copy, forms and operational documents")
    pdf.showPage()

    # Usage
    pdf.setFillColor(HexColor(WHITE))
    pdf.rect(0, 0, width, height, fill=1, stroke=0)
    pdf.setFillColor(HexColor(NAVY))
    pdf.setFont("Times-Roman", 27)
    pdf.drawString(46, height - 72, "Logo use")
    pdf.setFont("Helvetica", 10)
    pdf.drawString(46, height - 105, "Use the primary lockup whenever space permits.")
    pdf.drawImage(
        ImageReader(str(assets["dark"])),
        46,
        height - 220,
        width=340,
        height=68,
        mask="auto",
    )
    pdf.setFillColor(HexColor(NAVY))
    pdf.roundRect(46, height - 415, width - 92, 145, 8, fill=1, stroke=0)
    pdf.drawImage(
        ImageReader(str(assets["light"])),
        72,
        height - 375,
        width=330,
        height=66,
        mask="auto",
    )
    pdf.setFillColor(HexColor("#176762"))
    pdf.setFont("Helvetica-Bold", 9)
    pdf.drawString(46, height - 470, "COMPACT MARK")
    pdf.drawImage(
        ImageReader(str(assets["monogram"])),
        46,
        height - 620,
        width=108,
        height=108,
        mask="auto",
    )
    pdf.setFillColor(HexColor(NAVY))
    pdf.setFont("Helvetica", 10)
    text = pdf.beginText(178, height - 535)
    text.setLeading(15)
    text.textLines(
        "Use the monogram for favicons, profile images and\n"
        "small-format applications. Never combine it with the\n"
        "old gradient DB icon. Preserve clear space equal to\n"
        "one quarter of the mark's diameter."
    )
    pdf.drawText(text)
    pdf.showPage()

    # Applications and naming
    pdf.setFillColor(HexColor(IVORY))
    pdf.rect(0, 0, width, height, fill=1, stroke=0)
    pdf.setFillColor(HexColor(NAVY))
    pdf.setFont("Times-Roman", 27)
    pdf.drawString(46, height - 72, "Applications and naming")
    pdf.setFont("Helvetica-Bold", 9)
    pdf.setFillColor(HexColor("#176762"))
    pdf.drawString(46, height - 120, "PUBLIC BRAND")
    pdf.setFillColor(HexColor(NAVY))
    pdf.setFont("Helvetica", 11)
    pdf.drawString(46, height - 148, "DeepBridge Advisory")
    pdf.setFont("Helvetica", 9.5)
    text = pdf.beginText(46, height - 178)
    text.setLeading(15)
    text.textLines(
        "Use on the website, proposals, presentations, email signatures,\n"
        "client and consultant communications, and social profiles."
    )
    pdf.drawText(text)
    pdf.setFont("Helvetica-Bold", 9)
    pdf.setFillColor(HexColor("#176762"))
    pdf.drawString(46, height - 250, "LEGAL ENTITY")
    pdf.setFillColor(HexColor(NAVY))
    pdf.setFont("Helvetica", 11)
    pdf.drawString(46, height - 278, "DUSTDEEP LTD trading as DeepBridge Advisory")
    pdf.setFont("Helvetica", 9.5)
    text = pdf.beginText(46, height - 308)
    text.setLeading(15)
    text.textLines(
        "Use in contracts, legal notices, invoices and formal footers where\n"
        "the contracting entity must be clear. Do not use DUSTDEEP as the\n"
        "headline-facing advisory brand."
    )
    pdf.drawText(text)
    pdf.setStrokeColor(HexColor(LINE))
    pdf.line(46, height - 390, width - 46, height - 390)
    pdf.setFillColor(HexColor(NAVY))
    pdf.setFont("Helvetica-Bold", 9)
    pdf.drawString(46, height - 430, "DOCUMENT FOOTER")
    pdf.setFont("Helvetica", 8.5)
    footer_example = (
        "DUSTDEEP LTD trading as DeepBridge Advisory | Company no. 16775578 | "
        "Registered office: Kemp House, 152-160 City Road, London, EC1V 2NX"
    )
    text = pdf.beginText(46, height - 460)
    text.setLeading(13)
    for line in (footer_example[:100], footer_example[100:]):
        text.textLine(line)
    pdf.drawText(text)
    pdf.setFillColor(HexColor(PALE_AQUA))
    pdf.roundRect(46, 100, width - 92, 118, 8, fill=1, stroke=0)
    pdf.setFillColor(HexColor("#176762"))
    pdf.setFont("Helvetica-Bold", 9)
    pdf.drawString(64, 190, "CONSISTENCY CHECK")
    pdf.setFillColor(HexColor(NAVY))
    pdf.setFont("Helvetica", 9)
    text = pdf.beginText(64, 168)
    text.setLeading(14)
    text.textLines(
        "Use only the supplied logo assets. Keep navy, aqua and ivory dominant.\n"
        "Avoid gradients, glossy effects, generic stock imagery and competing\n"
        "blue brand marks. Keep the tone clear, selective and delivery-led."
    )
    pdf.drawText(text)

    pdf.save()
    return path


def make_digital_assets() -> None:
    signature_html = f"""<!doctype html>
<html lang="en">
  <body style="margin:0;padding:0;font-family:Arial,Helvetica,sans-serif;color:{NAVY};">
    <table role="presentation" cellpadding="0" cellspacing="0" style="border-collapse:collapse;">
      <tr>
        <td style="padding-right:18px;vertical-align:top;border-right:2px solid {AQUA};">
          <img src="https://www.deepbridgeadvisory.co.uk/brand/deepbridge-logo-dark.png"
               width="210" alt="DeepBridge Advisory" style="display:block;border:0;">
        </td>
        <td style="padding-left:18px;vertical-align:top;">
          <div style="font-size:15px;font-weight:700;line-height:1.3;">[Name]</div>
          <div style="font-size:12px;color:{MUTED};line-height:1.6;">[Title]</div>
          <div style="font-size:12px;line-height:1.6;">
            <a href="mailto:{COMPANY['email']}" style="color:{NAVY};text-decoration:none;">{COMPANY['email']}</a><br>
            <a href="https://{COMPANY['website']}" style="color:{NAVY};text-decoration:none;">{COMPANY['website']}</a><br>
            <a href="https://{COMPANY['linkedin']}" style="color:#176762;text-decoration:none;">LinkedIn</a>
          </div>
        </td>
      </tr>
      <tr>
        <td colspan="2" style="padding-top:12px;font-size:9px;color:{MUTED};line-height:1.45;">
          {COMPANY['legal_name']} trading as {COMPANY['brand']}. Registered in {COMPANY['registered_in']}.
          Company no. {COMPANY['company_number']}. Registered office: {COMPANY['registered_office']}.
        </td>
      </tr>
    </table>
  </body>
</html>
"""
    (DIGITAL_DIR / "DeepBridge-Advisory-Email-Signature.html").write_text(
        signature_html,
        encoding="utf-8",
    )
    (DIGITAL_DIR / "DeepBridge-Advisory-Email-Signature.txt").write_text(
        "[Name]\n"
        "[Title] | DeepBridge Advisory\n"
        f"{COMPANY['email']} | {COMPANY['website']}\n"
        f"{COMPANY['legal_name']} trading as {COMPANY['brand']} | "
        f"Company no. {COMPANY['company_number']}\n",
        encoding="utf-8",
    )
    brand_tokens = {
        "brand": COMPANY["brand"],
        "legal_entity": COMPANY["legal_name"],
        "colors": {
            "deep_navy": NAVY,
            "navy_soft": NAVY_SOFT,
            "bridge_aqua": AQUA,
            "soft_aqua": AQUA_SOFT,
            "advisory_ivory": IVORY,
            "slate": MUTED,
        },
        "typography": {
            "display": "Georgia",
            "interface_and_documents": "Arial",
        },
        "website": f"https://{COMPANY['website']}",
    }
    (DIGITAL_DIR / "DeepBridge-Advisory-Brand-Tokens.json").write_text(
        json.dumps(brand_tokens, indent=2),
        encoding="utf-8",
    )


def make_readme() -> None:
    readme = f"""DEEPBRIDGE ADVISORY BRAND PACK

This pack standardises the public DeepBridge Advisory identity while keeping
{COMPANY['legal_name']} visible where the contracting entity must be clear.

01 Logos
  - Primary Dark: for white and light backgrounds
  - Primary Light: for navy and dark backgrounds
  - Monogram: for profile images, favicons and compact placements

02 Letterhead
  - Editable Word letterhead
  - Print-ready PDF letterhead

03 Contract Templates
  - Branded agreement shell for counsel-approved contractual wording

04 Onboarding
  - Consultant onboarding template
  - Client programme kickoff template

05 Digital
  - HTML and plain-text email signatures
  - JSON brand tokens for designers and developers

06 Brand Guidelines
  - Concise PDF covering logo, colour, typography and legal naming

USAGE
1. Replace all bracketed placeholder text before issue.
2. Use DeepBridge Advisory as the public-facing brand.
3. Use "{COMPANY['legal_name']} trading as {COMPANY['brand']}" in contracts,
   invoices, legal notices and formal document footers.
4. Contractual and compliance wording should be reviewed by appropriately
   qualified advisers. The included contract file is a visual shell, not legal advice.
5. Do not use the previous glossy blue-gradient DB icon.

CONTACT
{COMPANY['email']}
{COMPANY['website']}
"""
    (PACK_ROOT / "README.txt").write_text(readme, encoding="utf-8")


def zip_pack() -> None:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(PACK_ROOT.rglob("*")):
            if path.is_file():
                archive.write(path, Path(PACK_ROOT.name) / path.relative_to(PACK_ROOT))


def main() -> None:
    reset_output()
    assets = make_logo_assets()
    make_letterhead(assets["dark"])
    make_contract_template(assets["dark"])
    make_consultant_onboarding(assets["dark"])
    make_client_kickoff(assets["dark"])
    make_brand_guide(assets)
    make_digital_assets()
    make_readme()
    zip_pack()
    print(PACK_ROOT)
    print(ZIP_PATH)


if __name__ == "__main__":
    main()
